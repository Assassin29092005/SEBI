"""Assembly tests: both output targets round-trip through python-docx.

Synthetic checklist + GeneratedSections only — no real schema dependency, so
these tests pin the assembly contract in isolation (offline, no API keys).
"""

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument

from app.assemble.docx_builder import (
    CONTRADICTION_NOTICE,
    DRAFT_NOTICE,
    LEGAL_DISCLAIMER,
    assemble,
    format_inr_paise,
)
from app.facts import Fact, FactStore, Provenance, SourceKind
from app.generate.sections import Citation, GeneratedSection
from app.schema.models import (
    Checklist,
    ChecklistEntry,
    ChecklistHeader,
    OutputTarget,
    Role,
    Severity,
)

DRHP_ONLY_TITLE = "History of equity share capital"
OVERVIEW_FACT_ID = "fact-incorporation-0001"
OVERVIEW_TEXT = (
    "Sunrise Agrotech Ltd was incorporated on 12 May 2015 in Pune, Maharashtra. "
    "The company mills and packages agro commodities. "
    "[REQUIRES INPUT: incorporation_certificate_number — promoter can provide this] "
    "It operates two processing facilities."
)
CAPITAL_TEXT = "The issuer has made three allotments of equity shares since incorporation."


def _checklist() -> Checklist:
    header = ChecklistHeader(
        regulation="SEBI ICDR Regulations, 2018 — Chapter IX",
        amended_through="2026-03-21",
        schema_version="test-0.0.1",
        reviewed_by_human=True,
    )
    entries = [
        ChecklistEntry(
            id="general.company_overview",
            clause_ref="ICDR Sch. VI Part A (as applied by Ch. IX), para (6)",
            section="General Information",
            title="Company overview",
            description="Overview of the issuer's business and incorporation history.",
            required_facts=["company_name", "incorporation_date"],
            responsible_role=Role.PROMOTER,
            severity=Severity.BLOCKER,
            output_targets=[OutputTarget.DRHP, OutputTarget.ABRIDGED],
        ),
        ChecklistEntry(
            id="capital_structure.share_capital_history",
            clause_ref="ICDR Sch. VI Part A, para (9)",
            section="Capital Structure",
            title=DRHP_ONLY_TITLE,
            description="Build-up of share capital since incorporation.",
            required_facts=["share_allotments"],
            responsible_role=Role.PROMOTER,
            severity=Severity.BLOCKER,
            output_targets=[OutputTarget.DRHP],
        ),
        ChecklistEntry(
            id="risk.risk_factors",
            clause_ref="ICDR Sch. VI Part A, para (5)",
            section="Risk Factors",
            title="Internal risk factors",
            description="Risks specific to the issuer and its business.",
            required_facts=["risk_factor_list"],
            responsible_role=Role.PROMOTER,
            severity=Severity.MATERIAL,
            output_targets=[OutputTarget.DRHP, OutputTarget.ABRIDGED],
        ),
    ]
    return Checklist(header=header, entries=entries)


def _sections() -> list[GeneratedSection]:
    first_sentence_end = OVERVIEW_TEXT.index(".") + 1
    return [
        GeneratedSection(
            entry_id="general.company_overview",
            section="General Information",
            text=OVERVIEW_TEXT,
            citations=[Citation(fact_id=OVERVIEW_FACT_ID, text_span=(0, first_sentence_end))],
            missing_facts=["incorporation_certificate_number"],
        ),
        GeneratedSection(
            entry_id="capital_structure.share_capital_history",
            section="Capital Structure",
            text=CAPITAL_TEXT,
            citations=[Citation(fact_id="fact-allotments-0002", text_span=(0, len(CAPITAL_TEXT)))],
            missing_facts=[],
        ),
        # risk.risk_factors deliberately has NO generated section → gap paragraph
    ]


def _assemble_both(tmp_path: Path) -> tuple[DocxDocument, DocxDocument]:
    checklist = _checklist()
    sections = _sections()
    drhp_path = assemble(checklist, sections, OutputTarget.DRHP, tmp_path / "drhp.docx")
    abridged_path = assemble(
        checklist, sections, OutputTarget.ABRIDGED, tmp_path / "abridged.docx"
    )
    assert drhp_path.exists() and abridged_path.exists()
    return Document(str(drhp_path)), Document(str(abridged_path))


def _paragraph_texts(doc: DocxDocument) -> list[str]:
    return [p.text for p in doc.paragraphs]


def test_draft_notice_on_both_covers(tmp_path: Path) -> None:
    drhp, abridged = _assemble_both(tmp_path)
    for doc in (drhp, abridged):
        assert any(DRAFT_NOTICE in text for text in _paragraph_texts(doc))


def test_titles_and_regulation_line(tmp_path: Path) -> None:
    drhp, abridged = _assemble_both(tmp_path)
    assert "Draft Red Herring Prospectus" in _paragraph_texts(drhp)
    assert "Draft Abridged Prospectus — Schedule VI Part E" in _paragraph_texts(abridged)
    for doc in (drhp, abridged):
        joined = "\n".join(_paragraph_texts(doc))
        assert "2026-03-21" in joined
        assert "test-0.0.1" in joined


def test_requires_input_runs_are_bold(tmp_path: Path) -> None:
    drhp, _ = _assemble_both(tmp_path)
    marker_runs = [
        run
        for paragraph in drhp.paragraphs
        for run in paragraph.runs
        if "[REQUIRES INPUT" in run.text
    ]
    assert marker_runs, "expected at least one [REQUIRES INPUT] run"
    for run in marker_runs:
        assert run.bold is True
        assert run.font.color.rgb is not None  # visually distinct (red)


def test_abridged_excludes_drhp_only_entry(tmp_path: Path) -> None:
    drhp, abridged = _assemble_both(tmp_path)
    assert any(DRHP_ONLY_TITLE in text for text in _paragraph_texts(drhp))
    assert not any(DRHP_ONLY_TITLE in text for text in _paragraph_texts(abridged))
    assert not any("Capital Structure" in text for text in _paragraph_texts(abridged))


def test_sources_list_and_superscript_markers(tmp_path: Path) -> None:
    drhp, _ = _assemble_both(tmp_path)
    texts = _paragraph_texts(drhp)
    assert any(text == "Sources" for text in texts)
    assert any(f"[1] {OVERVIEW_FACT_ID}" in text for text in texts)
    superscripts = [
        run
        for paragraph in drhp.paragraphs
        for run in paragraph.runs
        if run.font.superscript and run.text == "1"
    ]
    assert superscripts, "expected a superscript citation marker '1'"


def test_gap_paragraph_for_entry_without_generated_section(tmp_path: Path) -> None:
    drhp, abridged = _assemble_both(tmp_path)
    expected = "[REQUIRES INPUT: risk_factor_list — promoter can provide this]"
    for doc in (drhp, abridged):
        assert any(expected in text for text in _paragraph_texts(doc))


def test_toc_lists_target_sections(tmp_path: Path) -> None:
    drhp, abridged = _assemble_both(tmp_path)
    drhp_texts = _paragraph_texts(drhp)
    assert "Table of Contents" in drhp_texts
    assert "2. Capital Structure" in drhp_texts
    abridged_texts = _paragraph_texts(abridged)
    assert "1. General Information" in abridged_texts
    assert "2. Risk Factors" in abridged_texts


def test_format_inr_paise() -> None:
    assert format_inr_paise(25 * 10**9) == "₹25.00 crore"       # ₹25 crore
    assert format_inr_paise(15_50_00_000_00) == "₹15.50 crore"  # ₹15.5 crore
    assert format_inr_paise(25 * 10**6) == "₹2.50 lakh"         # ₹2.5 lakh
    assert format_inr_paise(1234567) == "₹12,345.67"
    assert format_inr_paise(123456789 * 100) == "₹12.35 crore"
    assert format_inr_paise(0) == "₹0.00"
    assert format_inr_paise(-150 * 10**9) == "-₹150.00 crore"


def test_issue_size_on_cover_uses_display_formatting(tmp_path: Path) -> None:
    path = assemble(
        _checklist(),
        _sections(),
        OutputTarget.DRHP,
        tmp_path / "drhp_sized.docx",
        issue_size_paise=24_50_00_000_00,  # ₹24.50 crore, arrives as paise
    )
    doc = Document(str(path))
    assert any("Indicative issue size: ₹24.50 crore" in t for t in _paragraph_texts(doc))


def _issue_size_fact(value: int, kind: SourceKind, detail: str) -> Fact:
    return Fact(
        key="issue_size_paise",
        value=value,
        provenance=Provenance(kind=kind, detail=detail),
        supplied_by=Role.PROMOTER,
    )


def _store_with_confirmed(*facts: Fact) -> FactStore:
    store = FactStore()
    for fact in facts:
        store.add(fact)
        store.confirm(fact.fact_id)
    return store


def test_store_single_confirmed_issue_size_renders_one_plain_line(tmp_path: Path) -> None:
    store = _store_with_confirmed(
        _issue_size_fact(24_50_00_000_00, SourceKind.WIZARD, "wizard: issue_size_paise"),
    )
    path = assemble(
        _checklist(), _sections(), OutputTarget.DRHP, tmp_path / "drhp_store.docx", store=store
    )
    texts = _paragraph_texts(Document(str(path)))
    size_lines = [t for t in texts if "Indicative issue size" in t]
    assert size_lines == ["Indicative issue size: ₹24.50 crore"]
    assert not any(CONTRADICTION_NOTICE in t for t in texts)


def test_store_contradiction_renders_both_values_and_bold_red_callout(tmp_path: Path) -> None:
    store = _store_with_confirmed(
        _issue_size_fact(24_50_00_000_00, SourceKind.WIZARD, "wizard: issue_size_paise"),
        _issue_size_fact(26_00_00_000_00, SourceKind.DOCUMENT, "bank_sanction_letter.txt p.1"),
    )
    path = assemble(
        _checklist(), _sections(), OutputTarget.DRHP, tmp_path / "drhp_contra.docx", store=store
    )
    doc = Document(str(path))
    texts = _paragraph_texts(doc)
    assert any(
        "Indicative issue size: ₹24.50 crore (wizard: issue_size_paise)" in t for t in texts
    )
    assert any(
        "Indicative issue size: ₹26.00 crore (bank_sanction_letter.txt p.1)" in t for t in texts
    )
    callout_runs = [
        run
        for paragraph in doc.paragraphs
        for run in paragraph.runs
        if CONTRADICTION_NOTICE.startswith(run.text) and run.text
    ]
    assert callout_runs, "expected the contradiction callout on the cover"
    for run in callout_runs:
        assert run.bold is True
        assert run.font.color.rgb is not None  # visually distinct (red)


def test_store_with_no_confirmed_issue_size_omits_the_line(tmp_path: Path) -> None:
    store = FactStore()
    store.add(  # present but never confirmed — must not leak onto the cover
        _issue_size_fact(24_50_00_000_00, SourceKind.WIZARD, "wizard: issue_size_paise")
    )
    path = assemble(
        _checklist(), _sections(), OutputTarget.DRHP, tmp_path / "drhp_empty.docx", store=store
    )
    texts = _paragraph_texts(Document(str(path)))
    assert not any("Indicative issue size" in t for t in texts)
    assert not any(CONTRADICTION_NOTICE in t for t in texts)


def test_legal_disclaimer_bold_on_both_covers(tmp_path: Path) -> None:
    drhp, abridged = _assemble_both(tmp_path)
    for doc in (drhp, abridged):
        texts = _paragraph_texts(doc)
        assert any(LEGAL_DISCLAIMER in t for t in texts)
        assert any(DRAFT_NOTICE in t for t in texts)  # existing notice kept alongside
        disclaimer_runs = [
            run
            for paragraph in doc.paragraphs
            for run in paragraph.runs
            if run.text and LEGAL_DISCLAIMER.startswith(run.text)
        ]
        assert disclaimer_runs, "expected the legal disclaimer on the cover"
        assert all(run.bold is True for run in disclaimer_runs)
