"""Document assembly via python-docx: the draft DRHP and the draft abridged prospectus.

RESOLVED (2026-07-03): the ICDR (Amendment) Regulations, 2026 (notified
2026-03-16, effective 2026-03-21) settle the second deliverable in favour of
the **draft abridged prospectus per Schedule VI Part E** — Reg. 246(3) requires
it to accompany the draft offer document filed with the SME exchange, and the
former "Offer Document Summary" (Sch. VI Part A para (4)) is omitted. See
``data/regulation/MANIFEST.md``. ``OutputTarget.ABRIDGED`` therefore assembles
the Part E draft abridged prospectus.

Layout rules implemented here (display layer only — no facts originate here):

- Cover page with the pinned regulation version from the checklist header and a
  prominent draft notice: the tool never produces a filing-ready document;
  merchant banker due diligence and certification come first.
- Static table-of-contents page (a live Word TOC field is a nice-to-have;
  production note: insert a ``TOC`` field and require a field refresh on open).
- One Heading 1 per schema section group, one Heading 2 per checklist entry
  with its clause_ref in small italics beneath it.
- ``[REQUIRES INPUT: ...]`` markers render as bold red runs so honest blanks
  are impossible to miss.
- Citations render as superscript numeric markers at each citation's end
  offset, with a per-entry "Sources" list mapping marker number → fact id, so
  every sentence stays traceable inside the document itself.
- Monetary values arrive as paise integers; lakh/crore formatting happens here
  and only here (see :func:`format_inr_paise`).
"""

from __future__ import annotations

import re
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Length, Pt, RGBColor
from docx.text.paragraph import Paragraph

from app.facts import FactStore
from app.generate.sections import GeneratedSection, requires_input_marker
from app.schema.models import Checklist, ChecklistEntry, ChecklistHeader, OutputTarget

DRAFT_NOTICE = (
    "DRAFT — NOT FOR FILING. Pending merchant banker due diligence and certification."
)

LEGAL_DISCLAIMER = (
    "This document is a computer-assisted draft and not legal advice. It may be filed "
    "only after due diligence and certification by a SEBI-registered merchant banker "
    "(lead manager)."
)

CONTRADICTION_NOTICE = (
    "CONTRADICTION DETECTED: confirmed sources disagree on the issue size — "
    "see the validation report before certification."
)

_TITLES: dict[OutputTarget, str] = {
    OutputTarget.DRHP: "Draft Red Herring Prospectus",
    OutputTarget.ABRIDGED: "Draft Abridged Prospectus — Schedule VI Part E",
}

_REQUIRES_INPUT_RE = re.compile(r"\[REQUIRES INPUT:[^\]]*\]")
_RED = RGBColor(0xC0, 0x00, 0x00)
_SMALL = Pt(8)

_PAISE_PER_CRORE = 10**9  # 1 crore rupees = 10^7 rupees = 10^9 paise
_PAISE_PER_LAKH = 10**7   # 1 lakh rupees  = 10^5 rupees = 10^7 paise


def format_inr_paise(paise: int) -> str:
    """Display-layer helper: paise integer → Indian-format INR string.

    ≥ 1 crore renders as "₹X.XX crore", ≥ 1 lakh as "₹X.XX lakh", below that
    as rupees with Indian digit grouping ("₹12,34,567.89"). Integer/Decimal
    arithmetic throughout — never floats.
    """
    sign = "-" if paise < 0 else ""
    magnitude = abs(paise)
    if magnitude >= _PAISE_PER_CRORE:
        value = (Decimal(magnitude) / _PAISE_PER_CRORE).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        return f"{sign}₹{value} crore"
    if magnitude >= _PAISE_PER_LAKH:
        value = (Decimal(magnitude) / _PAISE_PER_LAKH).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        return f"{sign}₹{value} lakh"
    rupees, remainder = divmod(magnitude, 100)
    return f"{sign}₹{_indian_group(rupees)}.{remainder:02d}"


def _indian_group(n: int) -> str:
    """Indian digit grouping: last three digits, then groups of two (12,34,567)."""
    s = str(n)
    if len(s) <= 3:
        return s
    head, tail = s[:-3], s[-3:]
    groups: list[str] = []
    while len(head) > 2:
        groups.append(head[-2:])
        head = head[:-2]
    if head:
        groups.append(head)
    return ",".join(reversed(groups)) + "," + tail


def assemble(
    checklist: Checklist,
    sections: list[GeneratedSection],
    target: OutputTarget,
    out_path: Path,
    issue_size_paise: int | None = None,
    store: FactStore | None = None,
) -> Path:
    """Assemble the formatted document for one output target.

    Checklist entries are filtered by ``target in entry.output_targets``, kept
    in YAML order, and grouped by ``entry.section``. Entries with no matching
    :class:`GeneratedSection` render as explicit gap paragraphs — honest blanks
    beat confident hallucination.

    ``issue_size_paise`` (optional) surfaces an indicative issue size on the
    cover page, formatted lakh/crore at this display layer only.

    ``store`` (optional) makes the cover page pull live confirmed
    ``issue_size_paise`` facts instead: more than one distinct confirmed value
    renders every value with its provenance plus a contradiction callout, so a
    cross-source disagreement is visible in the exported artefact itself.
    """
    entries = [e for e in checklist.entries if target in e.output_targets]
    grouped: dict[str, list[ChecklistEntry]] = {}
    for entry in entries:
        grouped.setdefault(entry.section, []).append(entry)
    generated_by_id: dict[str, GeneratedSection] = {}
    for gen in sections:
        generated_by_id.setdefault(gen.entry_id, gen)

    issue_size_lines, contradiction = _issue_size_lines(store, issue_size_paise)
    doc = Document()
    _add_cover_page(doc, checklist.header, target, issue_size_lines, contradiction)
    doc.add_page_break()
    _add_toc_page(doc, list(grouped))
    doc.add_page_break()
    for section_name, group in grouped.items():
        doc.add_heading(section_name, level=1)
        for entry in group:
            _add_entry(doc, entry, generated_by_id.get(entry.id))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _issue_size_lines(
    store: FactStore | None,
    issue_size_paise: int | None,
) -> tuple[list[str], bool]:
    """Cover-page issue-size lines and whether confirmed sources contradict.

    With a store, live confirmed ``issue_size_paise`` facts win. Exactly one
    distinct value keeps the plain single-line rendering; more than one
    distinct value renders each value with its provenance detail so the
    contradiction is a visible artefact in the exported document. Zero
    confirmed facts (or no store) fall back to the explicit
    ``issue_size_paise`` parameter — the original behaviour.
    """
    if store is not None:
        details_by_value: dict[int, list[str]] = {}
        for fact in store.confirmed_by_key("issue_size_paise"):
            if isinstance(fact.value, int) and not isinstance(fact.value, bool):
                details_by_value.setdefault(fact.value, []).append(fact.provenance.detail)
        if len(details_by_value) == 1:
            (value,) = details_by_value
            return [f"Indicative issue size: {format_inr_paise(value)}"], False
        if len(details_by_value) > 1:
            return [
                f"Indicative issue size: {format_inr_paise(value)} ({'; '.join(details)})"
                for value, details in details_by_value.items()
            ], True
    if issue_size_paise is not None:
        return [f"Indicative issue size: {format_inr_paise(issue_size_paise)}"], False
    return [], False


def _add_cover_page(
    doc: DocxDocument,
    header: ChecklistHeader,
    target: OutputTarget,
    issue_size_lines: list[str],
    contradiction: bool,
) -> None:
    title = doc.add_heading(_TITLES[target], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if target is OutputTarget.ABRIDGED:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_styled_run(
            subtitle,
            "Draft abridged prospectus accompanying the draft offer document (ICDR Reg. 246(3))",
            italic=True,
        )
    regulation = doc.add_paragraph()
    regulation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(
        regulation,
        f"{header.regulation} — as amended through {header.amended_through} "
        f"(checklist schema {header.schema_version})",
    )
    for line in issue_size_lines:
        size_p = doc.add_paragraph()
        size_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_styled_run(size_p, line)
    if contradiction:
        warning = doc.add_paragraph()
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_styled_run(warning, CONTRADICTION_NOTICE, bold=True, color=_RED)
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(notice, DRAFT_NOTICE, bold=True, color=_RED, size=Pt(14))
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(disclaimer, LEGAL_DISCLAIMER, bold=True)


def _add_toc_page(doc: DocxDocument, section_names: list[str]) -> None:
    doc.add_heading("Table of Contents", level=1)
    for number, name in enumerate(section_names, start=1):
        doc.add_paragraph(f"{number}. {name}")


def _add_entry(doc: DocxDocument, entry: ChecklistEntry, gen: GeneratedSection | None) -> None:
    doc.add_heading(entry.title, level=2)
    clause_p = doc.add_paragraph()
    _add_styled_run(clause_p, entry.clause_ref, italic=True, size=_SMALL)
    if gen is None:
        _add_gap_paragraph(doc, entry)
        return
    _add_generated_text(doc, gen)
    _add_sources_list(doc, gen)


def _add_gap_paragraph(doc: DocxDocument, entry: ChecklistEntry) -> None:
    """No generated content for this requirement — render the gap explicitly."""
    gap_p = doc.add_paragraph()
    _add_styled_run(
        gap_p,
        f"Requirement not yet drafted: {entry.description.strip()} ",
        italic=True,
    )
    facts = ", ".join(entry.required_facts) if entry.required_facts else entry.id
    _add_styled_run(
        gap_p,
        requires_input_marker(facts, str(entry.responsible_role)),
        bold=True,
        color=_RED,
    )


def _add_generated_text(doc: DocxDocument, gen: GeneratedSection) -> None:
    """Render section text with bold-red [REQUIRES INPUT] runs and superscript citations."""
    text = gen.text
    paragraph = doc.add_paragraph()
    position = 0
    for offset, number, _fact_id in _numbered_citations(gen):
        _add_marked_runs(paragraph, text[position:offset])
        _add_styled_run(paragraph, str(number), superscript=True)
        position = offset
    _add_marked_runs(paragraph, text[position:])


def _add_sources_list(doc: DocxDocument, gen: GeneratedSection) -> None:
    """Per-entry Sources list: superscript marker number → fact id (traceability)."""
    numbered = _numbered_citations(gen)
    if not numbered:
        return
    label = doc.add_paragraph()
    _add_styled_run(label, "Sources", bold=True, size=_SMALL)
    for _offset, number, fact_id in numbered:
        item = doc.add_paragraph()
        _add_styled_run(item, f"[{number}] {_short_fact_id(fact_id)}", size=_SMALL)


def _numbered_citations(gen: GeneratedSection) -> list[tuple[int, int, str]]:
    """(end_offset, marker_number, fact_id) sorted by end offset, numbered 1..n.

    Offsets are clamped to the text bounds; an offset falling inside a
    [REQUIRES INPUT] marker is snapped to the marker's end so the marker is
    never split (which would break its bold-red rendering).
    """
    text = gen.text
    marker_spans = [(m.start(), m.end()) for m in _REQUIRES_INPUT_RE.finditer(text)]

    def snap(offset: int) -> int:
        offset = max(0, min(offset, len(text)))
        for start, end in marker_spans:
            if start < offset < end:
                return end
        return offset

    ordered = sorted(gen.citations, key=lambda c: c.text_span[1])
    return [
        (snap(citation.text_span[1]), number, citation.fact_id)
        for number, citation in enumerate(ordered, start=1)
    ]


def _short_fact_id(fact_id: str) -> str:
    return fact_id if len(fact_id) <= 40 else fact_id[:37] + "..."


def _add_marked_runs(paragraph: Paragraph, text: str) -> None:
    """Split text around [REQUIRES INPUT: ...] markers; markers become bold red runs."""
    position = 0
    for match in _REQUIRES_INPUT_RE.finditer(text):
        _add_styled_run(paragraph, text[position : match.start()])
        _add_styled_run(paragraph, match.group(), bold=True, color=_RED)
        position = match.end()
    _add_styled_run(paragraph, text[position:])


def _add_styled_run(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    size: Length | None = None,
    superscript: bool = False,
) -> None:
    """Add styled run(s); embedded newlines become explicit line breaks."""
    for index, line in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        if not line:
            continue
        run = paragraph.add_run(line)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = size
        if superscript:
            run.font.superscript = True
